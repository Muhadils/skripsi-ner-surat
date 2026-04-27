#!/usr/bin/env python3
"""
TEXT GENERATION SYSTEM using IndoBART
Skripsi Nurfadilah Rahman - Pembuatan Teks Otomatis
"""

import json
import torch
import os
from datetime import datetime
from io import BytesIO
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

class SuratAutomatisGenerator:
    """
    Generate surat administrasi otomatis dari extracted entities
    """
    def __init__(self, model_name="indobenchmark/indobart"):
        print("Initializing IndoBART model...")
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            self.model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
            self.generator = pipeline(
                "text2text-generation",
                model=self.model,
                tokenizer=self.tokenizer,
                device="cpu"  # Use CPU for compatibility
            )
            print(f"✓ Model loaded: {model_name}")
        except Exception as e:
            print(f"⚠️  Warning: Could not load {model_name}: {e}")
            print("   Will use template-based generation instead")
            self.generator = None

    def extract_entities_from_text(self, text):
        """
        Extract entities dari teks surat
        Menggunakan regex patterns dan simple matching
        """
        entities = {
            "nama": self._extract_nama(text),
            "nik": self._extract_nik(text),
            "alamat": self._extract_alamat(text),
            "pekerjaan": self._extract_pekerjaan(text),
            "tanggal": self._extract_tanggal(text),
            "jabatan": self._extract_jabatan(text),
        }
        return entities

    def _extract_nama(self, text):
        """Extract nama orang dari teks"""
        # Pola: nama di antara kalimat
        patterns = [
            r'bernama\s+([A-Z][a-z\s]+)',
            r'(?:Bapak|Ibu|Tn\.?|Ny\.?)\s+([A-Z][a-z\s]+?)(?:\s+(?:dari|bin|binti|bin|alias|atau))',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "Nama Lengkap"

    def _extract_nik(self, text):
        """Extract NIK (16 digit nomor)"""
        match = re.search(r'\b(\d{16})\b', text)
        return match.group(1) if match else "0000000000000000"

    def _extract_alamat(self, text):
        """Extract alamat dari teks"""
        patterns = [
            r'(?:bertempat tinggal|bertempat|domisili|alamat)(?:\s+di)?\s+([A-Za-z\s,\.]+?)(?:\s+(?:Kel\.|Kec\.|Kab\.|yang|dengan))',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "Alamat, Kelurahan, Kecamatan, Kabupaten"

    def _extract_pekerjaan(self, text):
        """Extract pekerjaan/profesi"""
        patterns = [
            r'(?:pekerjaan|profesi|sebagai)\s+([A-Za-z\s]+?)(?:\s+(?:dengan|dari|yang))',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "Pekerjaan/Profesi"

    def _extract_tanggal(self, text):
        """Extract tanggal dari teks"""
        match = re.search(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', text)
        if match:
            return f"{match.group(1)} {match.group(2)} {match.group(3)}"
        return "01 Januari 2025"

    def _extract_jabatan(self, text):
        """Extract jabatan/posisi"""
        patterns = [
            r'(?:sebagai|jabatan)\s+([A-Za-z\s]+?)(?:\s+(?:dari|di|yang))',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "Jabatan"

    def generate_surat_keterangan_domisili(self, entities):
        """Generate Surat Keterangan Domisili"""
        template = f"""Surat Keterangan Domisili

    Yang bertanda tangan di bawah ini Lurah Tadokkong, Kecamatan Lembang, Kabupaten Pinrang, Provinsi Sulawesi Selatan, menerangkan bahwa:

Nama                    : {entities.get('nama', 'Nama Lengkap')}
    Nomor Induk Kependudukan : {entities.get('nik', '0000000000000000')}
    Jenis Kelamin           : [Laki-laki/Perempuan]
    Alamat                  : {entities.get('alamat', 'Alamat, Kelurahan, Kecamatan, Kabupaten')}
    Pekerjaan               : {entities.get('pekerjaan', 'Pekerjaan/Profesi')}

    Adalah benar penduduk yang bertempat tinggal di wilayah Kelurahan Tadokkong, Kecamatan Lembang, Kabupaten Pinrang, Provinsi Sulawesi Selatan.

Surat keterangan ini diberikan untuk keperluan: {entities.get('maksud_tujuan', '[Sebutkan Keperluan]')}

Demikian surat keterangan ini dibuat dengan sebenar-benarnya dan dapat dipergunakan sebagai mana mestinya.

Diberikan di: Tadokkong
Tanggal: {entities.get('tanggal', '01 Januari 2025')}

Lurah Tadokkong
[Tanda Tangan dan Stempel]
[Nama Lengkap Lurah]
NIP. [Nomor Induk Pegawai]"""
        return template

    def generate_surat_keterangan_usaha(self, entities):
        """Generate Surat Keterangan Usaha"""
        template = f"""Surat Keterangan Usaha

    Yang bertanda tangan di bawah ini Lurah Tadokkong, Kecamatan Lembang, Kabupaten Pinrang, menerangkan bahwa:

Nama                    : {entities.get('nama', 'Nama Lengkap')}
    Nomor Induk Kependudukan : {entities.get('nik', '0000000000000000')}
    Alamat                  : {entities.get('alamat', 'Alamat')}
    Pekerjaan/Usaha         : {entities.get('pekerjaan', 'Jenis Usaha')}

Surat keterangan ini diberikan untuk keperluan: {entities.get('maksud_tujuan', '[Sebutkan Keperluan]')}

Surat keterangan ini berlaku selama 1 (satu) tahun apabila diperlukan dapat diperpanjang kembali.

Demikian surat keterangan ini dibuat dengan sebenar-benarnya untuk dapat dipergunakan sebagaimana mestinya.

Diberikan di: Tadokkong
Tanggal: {entities.get('tanggal', '01 Januari 2025')}

Lurah Tadokkong
[Tanda Tangan dan Stempel]
[Nama Lengkap Lurah]
NIP. [Nomor Induk Pegawai]"""
        return template

    def generate_surat_permohonan_pengantar(self, entities):
        """Generate Surat Pengantar/Permohonan"""
        template = f"""Surat Pengantar

Kepada Yth. [Instansi Tujuan]
di Tempat

    Dengan hormat, yang bertanda tangan di bawah ini Lurah Tadokkong menerangkan bahwa:

Nama                    : {entities.get('nama', 'Nama Lengkap')}
    Nomor Induk Kependudukan : {entities.get('nik', '0000000000000000')}
    Jenis Kelamin           : [Laki-laki/Perempuan]
    Alamat                  : {entities.get('alamat', 'Alamat')}

    Adalah benar penduduk Kelurahan Tadokkong, Kecamatan Lembang, Kabupaten Pinrang, Sulawesi Selatan dan memerlukan surat ini untuk keperluan:

{entities.get('maksud_tujuan', '[Sebutkan Maksud dan Tujuan]')}

Oleh karena itu kami mohon bantuan dan kerja sama Bapak/Ibu dalam memberikan pelayanan kepada yang bersangkutan sebagaimana mestinya.

Atas perhatian dan bantuan Bapak/Ibu diucapkan terimakasih.

Diberikan di: Tadokkong
Tanggal: {entities.get('tanggal', '01 Januari 2025')}

Lurah Tadokkong
[Tanda Tangan dan Stempel]
[Nama Lengkap Lurah]
NIP. [Nomor Induk Pegawai]"""
        return template

    def save_to_docx(self, text, filename):
        """Save generated text to DOCX file"""
        doc = self.create_docx_document(text)
        doc.save(filename)
        print(f"✓ Saved to {filename}")

    def create_docx_bytes(self, text):
        """Build DOCX and return bytes for web download."""
        doc = self.create_docx_document(text)
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def create_docx_document(self, text):
        """Create polished DOCX format for surat administratif."""
        doc = Document()

        # Global page layout
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.9)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        lines = text.split("\n")
        signature_mode = False
        surat_title = None
        title_skipped = False

        for raw_line in lines:
            candidate = raw_line.strip()
            if candidate:
                surat_title = candidate
                break

        self._add_letterhead(doc, surat_title)

        for raw_line in lines:
            line = raw_line.strip()

            if not line:
                spacer = doc.add_paragraph("")
                spacer.paragraph_format.space_after = Pt(1)
                continue

            if not title_skipped:
                title_skipped = True
                continue

            if line.startswith("Diberikan di:") or line.startswith("Tanggal:"):
                p = doc.add_paragraph(line)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                p.paragraph_format.space_after = Pt(0)
                signature_mode = True
                continue

            if line in {"[Tanda Tangan dan Stempel]"}:
                spacer = doc.add_paragraph("")
                spacer.paragraph_format.space_after = Pt(4)
                line_block = doc.add_paragraph("______________________________")
                line_block.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                line_block.paragraph_format.space_after = Pt(2)
                continue

            if line in {"Lurah Tadokkong", "[Nama Lengkap Lurah]", "[Nomor Induk Pegawai]"}:
                p = doc.add_paragraph(line)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                p.paragraph_format.space_after = Pt(0)
                if line in {"Lurah Tadokkong", "[Nama Lengkap Lurah]"}:
                    p.runs[0].bold = True
                continue

            if self._is_biodata_line(line):
                label, value = [part.strip() for part in line.split(":", 1)]
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.1
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.tab_stops.add_tab_stop(Inches(2.75))
                p.add_run(self._format_biodata_label(label))
                p.add_run("\t: ")
                p.add_run(value)
                continue

            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.1
            p.paragraph_format.space_after = Pt(1)

            if signature_mode:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line.startswith("Kepada Yth") or line == "di Tempat":
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        return doc

    def _add_letterhead(self, doc, surat_title):
        """Add a simple formal letterhead to the document."""
        logo_path = self._find_logo_path()

        if logo_path:
            try:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                self._remove_table_borders(table)
                left_cell, right_cell = table.rows[0].cells
                left_cell.width = Cm(3.0)
                right_cell.width = Cm(13.0)

                logo_paragraph = left_cell.paragraphs[0]
                logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.15))

                header_paragraph = right_cell.paragraphs[0]
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                header_paragraph.paragraph_format.space_after = Pt(0)
                run1 = header_paragraph.add_run("PEMERINTAH KELURAHAN TADOKKONG")
                run1.bold = True
                run1.font.size = Pt(12)
                run2 = header_paragraph.add_run("\nKECAMATAN LEMBANG KABUPATEN PINRANG")
                run2.bold = True
                run2.font.size = Pt(12)
                run3 = header_paragraph.add_run("\nJl. Poros Tadokkong, Kecamatan Lembang, Kabupaten Pinrang")
                run3.font.size = Pt(10)

                line = doc.add_paragraph("=" * 60)
                line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                line.paragraph_format.space_after = Pt(6)
            except (UnrecognizedImageError, OSError, ValueError):
                # Fall back to text-only header when logo file is invalid/corrupt.
                logo_path = None

        if not logo_path:
            header_lines = [
                "PEMERINTAH KELURAHAN TADOKKONG",
                "KECAMATAN LEMBANG KABUPATEN PINRANG",
                "Jl. Poros Tadokkong, Kecamatan Lembang, Kabupaten Pinrang",
            ]

            for idx, header_line in enumerate(header_lines):
                p = doc.add_paragraph(header_line)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(0 if idx < len(header_lines) - 1 else 2)
                run = p.runs[0]
                run.bold = True if idx < 2 else False
                run.font.size = Pt(11 if idx < 2 else 9)

            line = doc.add_paragraph("=" * 60)
            line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            line.paragraph_format.space_after = Pt(3)

        title_code = self._generate_nomor_surat(surat_title)
        nomor = doc.add_paragraph(f"Nomor\t: {title_code}")
        nomor.paragraph_format.space_after = Pt(0)
        perihal = doc.add_paragraph(f"Perihal\t: {surat_title.upper() if surat_title else 'SURAT KETERANGAN'}")
        perihal.paragraph_format.space_after = Pt(4)

    def _generate_nomor_surat(self, surat_title):
        """Generate a simple formal letter number based on the title."""
        now = datetime.now()
        month_romawi = [
            "I", "II", "III", "IV", "V", "VI",
            "VII", "VIII", "IX", "X", "XI", "XII"
        ][now.month - 1]

        normalized = (surat_title or "SURAT").lower()
        if "domisili" in normalized:
            code = "DS"
        elif "usaha" in normalized:
            code = "US"
        elif "pengantar" in normalized:
            code = "PG"
        else:
            code = "SK"

        return f"140/{code}/TDK/{month_romawi}/{now.year}"

    def _is_biodata_line(self, line):
        """Detect identity rows that should be aligned with a colon."""
        if ":" not in line:
            return False
        label = line.split(":", 1)[0].strip().lower()
        known_labels = {
            "nama",
            "nomor induk kependudukan",
            "jenis kelamin",
            "alamat",
            "pekerjaan",
            "pekerjaan/usaha",
        }
        return label in known_labels

    def _format_biodata_label(self, label):
        """Normalize biodata labels so alignment stays consistent in DOCX."""
        target_width = 24
        return label.ljust(target_width)

    def _find_logo_path(self):
        """Locate the logo file from datasurat or common local paths."""
        candidate_paths = [
            "datasurat/Logo_Universitas_Muhammadiyah_Makassar_Resmi.jpg",
            "logo.png",
            "logo.jpg",
            "logo.jpeg",
            "logo_um.png",
            "logo_um.jpg",
            "logo_um.jpeg",
            "logo_unismuh.png",
            "logo_unismuh.jpg",
            "assets/logo.png",
            "assets/logo.jpg",
            "static/logo.png",
            "static/logo.jpg",
        ]

        for candidate in candidate_paths:
            if os.path.exists(candidate):
                return candidate
        return None

    def _remove_table_borders(self, table):
        """Remove visible borders from the header table so the letterhead looks clean."""
        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_el = tbl_borders.find(qn(f"w:{edge}"))
            if edge_el is None:
                edge_el = OxmlElement(f"w:{edge}")
                tbl_borders.append(edge_el)
            edge_el.set(qn("w:val"), "nil")
            edge_el.set(qn("w:sz"), "0")
            edge_el.set(qn("w:space"), "0")
            edge_el.set(qn("w:color"), "auto")

    def process_surat(self,surat_type, entities):
        """Process and generate specific surat type"""
        if surat_type == "domisili":
            return self.generate_surat_keterangan_domisili(entities)
        elif surat_type == "usaha":
            return self.generate_surat_keterangan_usaha(entities)
        elif surat_type == "pengantar":
            return self.generate_surat_permohonan_pengantar(entities)
        else:
            return "Tipe surat tidak dikenal"


def main():
    """Demo usage"""
    print("="*70)
    print("  SISTEM GENERASI SURAT ADMINISTRASI OTOMATIS")
    print("="*70)

    # Initialize generator
    generator = SuratAutomatisGenerator()

    # Example entities (biasanya dari NER extraction)
    example_entities = {
        "nama": "Nurfadilah Rahman",
        "nik": "7315076512800001",
        "alamat": "Tuppu, Kel. Tadokkong, Kec. Lembang, Kab. Pinrang",
        "pekerjaan": "Mahasiswa/Karyawan",
        "tanggal": "25 April 2026",
        "jabatan": "Operator",
    }

    # Generate different types of surats
    surat_types = ["domisili", "usaha", "pengantar"]

    for surat_type in surat_types:
        print(f"\n📄 Generating: Surat Keterangan {surat_type.upper()}")
        surat_text = generator.process_surat(surat_type, example_entities)

        # Print preview
        print("\n" + "="*70)
        print(surat_text[:500] + "...")
        print("="*70)

        # Save to DOCX
        output_file = f"./hasil_surat_{surat_type}.docx"
        generator.save_to_docx(surat_text, output_file)

    print("\n✅ Generation completed!")
    print("📦 Check hasil_surat_*.docx files for generated documents")


if __name__ == "__main__":
    main()
